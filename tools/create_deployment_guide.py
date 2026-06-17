from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "deployment-guide-auto-parts-crm.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(91, 108, 126)
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F8FD"
WHITE = "FFFFFF"
BORDER = "C9D8EA"


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "bottom": bottom,
        "start": start,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_paragraph(paragraph, before=0, after=6, line_spacing=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    style_paragraph(
        paragraph,
        before={1: 18, 2: 14, 3: 10}[level],
        after={1: 10, 2: 7, 3: 5}[level],
    )
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}[level],
        color=BLUE if level < 3 else DARK_BLUE,
        bold=True,
    )
    return paragraph


def add_body(doc, text, after=6, bold_prefix=None):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        prefix_run = paragraph.add_run(bold_prefix)
        set_run_font(prefix_run, size=11, color=NAVY, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(run, size=11, color=RGBColor(34, 47, 62))
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=RGBColor(34, 47, 62))
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    style_paragraph(paragraph, after=4)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=RGBColor(34, 47, 62))


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    style_paragraph(paragraph, after=4)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=RGBColor(34, 47, 62))


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, "D9E2EF")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7FAFD")
    set_cell_margins(cell, top=110, bottom=110, start=150, end=150)
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, after=0, line_spacing=1.12)
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=9.5, color=RGBColor(31, 42, 55))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, "B7D7F1")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EFF8FF")
    set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, after=0, line_spacing=1.2)
    title_run = paragraph.add_run(f"{title}: ")
    set_run_font(title_run, size=10.5, color=DARK_BLUE, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=RGBColor(34, 47, 62))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        style_paragraph(paragraph, after=0, line_spacing=1.15)
        run = paragraph.add_run(header)
        set_run_font(run, size=9.5, color=NAVY, bold=True)
    for row in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = row_cells[idx]
            set_cell_shading(cell, WHITE if len(table.rows) % 2 else PALE_BLUE)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            style_paragraph(paragraph, after=0, line_spacing=1.15)
            run = paragraph.add_run(value)
            set_run_font(run, size=9.3, color=RGBColor(34, 47, 62))
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Auto Parts CRM Deployment Guide")
    set_run_font(footer_run, size=8.5, color=MUTED)
    return doc


def add_cover(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run("DEPLOYMENT GUIDE")
    set_run_font(run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Auto Parts CRM")
    set_run_font(run, size=28, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("AWS deployment process using Amazon RDS PostgreSQL")
    set_run_font(run, size=14, color=MUTED)

    add_callout(
        doc,
        "Use this guide when deploying",
        "Follow the sections in order: create AWS database, prepare server, configure environment variables, run Prisma migrations, build both apps, start services, and configure Nginx/SSL.",
    )

    add_matrix(
        doc,
        ["Layer", "Recommended AWS service", "Purpose"],
        [
            ["Database", "Amazon RDS for PostgreSQL", "Managed PostgreSQL database for Prisma models."],
            ["Compute", "EC2 Ubuntu", "Runs NestJS backend and Next.js frontend Node servers."],
            ["Cache / Queue", "ElastiCache Redis, optional", "Redis cache and BullMQ support when enabled."],
            ["Proxy", "Nginx on EC2", "Routes frontend and API domains to local Node ports."],
            ["Process manager", "PM2", "Keeps backend and frontend alive after reboot/deploy."],
        ],
        [1.45, 2.15, 2.9],
    )

    add_body(
        doc,
        "Production ports: backend runs on 3000 and frontend runs on 3001 before Nginx maps them to public domains.",
    )


def add_sections(doc):
    add_heading(doc, "1. Create AWS RDS PostgreSQL", 1)
    for item in [
        "Open AWS Console, go to RDS > Databases > Create database.",
        "Select PostgreSQL and choose a Dev/Test or Production template.",
        "Use database name auto_parts_crm, then set a strong username and password.",
        "Keep the DB inside your VPC. Prefer private access instead of public access.",
        "Configure the DB security group to allow port 5432 only from the EC2 server security group.",
        "Enable automated backups and copy the RDS endpoint after creation.",
    ]:
        add_numbered(doc, item)
    add_code_block(
        doc,
        [
            'DATABASE_URL="postgresql://DB_USER:DB_PASSWORD@RDS_ENDPOINT:5432/auto_parts_crm?schema=public"',
        ],
    )

    add_heading(doc, "2. Prepare EC2 Ubuntu Server", 1)
    add_code_block(
        doc,
        [
            "sudo apt update",
            "sudo apt install -y git nginx",
            "curl -fsSL https://deb.nodesource.com/setup_22.x | sudo -E bash -",
            "sudo apt install -y nodejs",
            "sudo npm install -g pm2",
        ],
    )
    add_code_block(
        doc,
        [
            "git clone <your-repo-url>",
            "cd auto-parts-crm",
            "npm install --prefix backend",
            "npm install --prefix frontend",
        ],
    )

    add_heading(doc, "3. Backend Environment", 1)
    add_body(doc, "Create backend/.env with production values:")
    add_code_block(
        doc,
        [
            'DATABASE_URL="postgresql://DB_USER:DB_PASSWORD@RDS_ENDPOINT:5432/auto_parts_crm?schema=public"',
            'JWT_SECRET="use-a-long-secure-random-secret"',
            "JWT_EXPIRES_IN_SECONDS=86400",
            'REDIS_HOST="your-redis-host-or-127.0.0.1"',
            "REDIS_PORT=6379",
            'REDIS_USERNAME=""',
            'REDIS_PASSWORD=""',
            "REDIS_DB=0",
            "REDIS_CACHE_ENABLED=true",
            'REDIS_CACHE_KEY_PREFIX="auto-parts-crm:cache:"',
            "REDIS_TLS_ENABLED=false",
            "REDIS_CONNECT_TIMEOUT_MS=5000",
            "ORDERS_CACHE_TTL_SECONDS=60",
            'BULLMQ_PREFIX="auto-parts-crm"',
        ],
    )
    add_callout(
        doc,
        "Redis note",
        "If Redis is not deployed yet, set REDIS_CACHE_ENABLED=false. If BullMQ workers are required, deploy Redis through AWS ElastiCache or another managed Redis service.",
    )

    add_heading(doc, "4. Frontend Environment", 1)
    add_body(doc, "Create frontend/.env.production:")
    add_code_block(
        doc,
        [
            'NEXT_PUBLIC_APP_NAME="Auto Parts CRM"',
            'NEXT_PUBLIC_APP_URL="https://your-domain.com"',
            "NEXT_PUBLIC_API_TIMEOUT_MS=10000",
            "NEXT_PUBLIC_TOAST_DURATION_MS=5000",
            'BACKEND_API_URL="https://api.your-domain.com"',
            "BACKEND_API_TIMEOUT_MS=10000",
            "AUTH_COOKIE_MAX_AGE_SECONDS=86400",
        ],
    )

    add_heading(doc, "5. Run Prisma Production Migration", 1)
    add_body(
        doc,
        "Run production migrations against the RDS database. Use migrate deploy, not migrate dev, for deployment.",
    )
    add_code_block(
        doc,
        [
            "cd backend",
            "npm run prisma:generate",
            "npx prisma migrate deploy",
            "cd ..",
        ],
    )

    add_heading(doc, "6. Build Backend And Frontend", 1)
    add_body(doc, "On Linux EC2:")
    add_code_block(
        doc,
        [
            "npm run build --prefix backend",
            "npm run build --prefix frontend",
        ],
    )
    add_body(doc, "On Windows local machine:")
    add_code_block(
        doc,
        [
            "npm.cmd run build --prefix backend",
            "npm.cmd run build --prefix frontend",
        ],
    )

    add_heading(doc, "7. Start Production Apps With PM2", 1)
    add_code_block(
        doc,
        [
            'pm2 start "npm run start:prod --prefix backend" --name auto-parts-backend',
            'pm2 start "npm run start --prefix frontend" --name auto-parts-frontend',
            "pm2 save",
            "pm2 startup",
        ],
    )

    add_heading(doc, "8. Configure Nginx Reverse Proxy", 1)
    add_code_block(
        doc,
        [
            "server {",
            "    server_name your-domain.com;",
            "    location / {",
            "        proxy_pass http://127.0.0.1:3001;",
            "        proxy_http_version 1.1;",
            "        proxy_set_header Host $host;",
            "        proxy_set_header X-Real-IP $remote_addr;",
            "    }",
            "}",
            "",
            "server {",
            "    server_name api.your-domain.com;",
            "    location / {",
            "        proxy_pass http://127.0.0.1:3000;",
            "        proxy_http_version 1.1;",
            "        proxy_set_header Host $host;",
            "        proxy_set_header X-Real-IP $remote_addr;",
            "    }",
            "}",
        ],
    )
    add_code_block(
        doc,
        [
            "sudo apt install -y certbot python3-certbot-nginx",
            "sudo certbot --nginx",
        ],
    )

    add_heading(doc, "9. Deployment Runbook", 1)
    add_body(doc, "Use this sequence for future deployments:")
    add_code_block(
        doc,
        [
            "git pull",
            "npm install --prefix backend",
            "npm install --prefix frontend",
            "npm run prisma:generate --prefix backend",
            "cd backend && npx prisma migrate deploy && cd ..",
            "npm run build --prefix backend",
            "npm run build --prefix frontend",
            "pm2 restart auto-parts-backend",
            "pm2 restart auto-parts-frontend",
        ],
    )

    add_heading(doc, "10. Production Checklist", 1)
    checklist = [
        "RDS is private and port 5432 is not exposed publicly.",
        "EC2 security group allows HTTP/HTTPS and restricts SSH.",
        "Backend .env uses the RDS endpoint and a strong JWT_SECRET.",
        "Prisma migrate deploy has completed successfully.",
        "Backend build and frontend build both pass.",
        "PM2 shows both processes online.",
        "Nginx routes frontend and API domains correctly.",
        "SSL certificate is installed and auto-renewal is enabled.",
        "RDS automated backups are enabled.",
    ]
    for item in checklist:
        add_bullet(doc, item)

    add_heading(doc, "References", 1)
    refs = [
        "AWS RDS PostgreSQL documentation: https://aws.amazon.com/rds/postgresql/",
        "AWS RDS create database documentation: https://docs.aws.amazon.com/AmazonRDS/latest/UserGuide/USER_CreateDBInstance.html",
        "AWS RDS security groups documentation: https://docs.aws.amazon.com/AmazonRDS/latest/UserGuide/Overview.RDSSecurityGroups.html",
        "Prisma production migration documentation: https://www.prisma.io/docs/orm/prisma-client/deployment/deploy-database-changes-with-prisma-migrate",
        "Next.js deployment documentation: https://nextjs.org/docs/app/getting-started/deploying",
    ]
    for ref in refs:
        add_bullet(doc, ref)


def main():
    doc = setup_document()
    add_cover(doc)
    add_sections(doc)
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
